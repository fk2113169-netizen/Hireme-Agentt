import os
import fitz  # PyMuPDF
import docx

def extract_text_from_pdf(file_path_or_bytes) -> str:
    """Extracts all text from a PDF file."""
    text = ""
    # Check if input is bytes or path
    if isinstance(file_path_or_bytes, bytes):
        doc = fitz.open(stream=file_path_or_bytes, filetype="pdf")
    else:
        doc = fitz.open(file_path_or_bytes)
        
    for page in doc:
        text += page.get_text()
    doc.close()
    return text

def extract_text_from_docx(file_path_or_bytes) -> str:
    """Extracts all text from a Word document (.docx)."""
    # If bytes, we write to a temporary file or read as stream (docx.Document can read from bytes stream)
    import io
    if isinstance(file_path_or_bytes, bytes):
        stream = io.BytesIO(file_path_or_bytes)
        doc = docx.Document(stream)
    else:
        doc = docx.Document(file_path_or_bytes)
        
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return "\n".join(text)

def extract_cv_text(file_name: str, file_bytes: bytes) -> str:
    """Helper to extract text based on file extension."""
    ext = os.path.splitext(file_name)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")
